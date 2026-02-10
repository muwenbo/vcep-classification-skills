#!/usr/bin/env python3
"""
Read Word documents (.docx) and convert to markdown format.

Usage:
    python read_word.py <path_to_word_file> [--json]

Arguments:
    path_to_word_file    Path to the Word document to read
    --json               Output as JSON instead of markdown (optional)

Output:
    Markdown text with headings preserved (default) or JSON structure

Dependencies:
    python-docx (auto-installed if missing)
"""

import sys
import json
import argparse

def install_dependencies():
    """Install required packages if not present."""
    import subprocess
    try:
        from docx import Document
    except ImportError:
        print("Installing python-docx...", file=sys.stderr)
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "python-docx", "-q"],
            stderr=subprocess.DEVNULL
        )

def read_docx_file(file_path):
    """
    Read a .docx file and extract its contents.

    Args:
        file_path: Path to the Word document

    Returns:
        Dictionary with document structure and content
    """
    try:
        from docx import Document
    except ImportError:
        install_dependencies()
        from docx import Document

    doc = Document(file_path)
    result = {
        'paragraphs': [],
        'tables': [],
        'metadata': {}
    }

    # Extract core properties if available
    try:
        core_props = doc.core_properties
        result['metadata'] = {
            'title': core_props.title or '',
            'author': core_props.author or '',
            'created': str(core_props.created) if core_props.created else '',
            'modified': str(core_props.modified) if core_props.modified else ''
        }
    except Exception:
        pass

    # Extract paragraphs with style information
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style_name = para.style.name if para.style else "Normal"
            result['paragraphs'].append({
                'text': text,
                'style': style_name
            })

    # Extract tables
    for table_idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        result['tables'].append({
            'index': table_idx,
            'rows': len(table.rows),
            'cols': len(table.rows[0].cells) if table.rows else 0,
            'data': table_data
        })

    return result

def format_as_markdown(data):
    """Convert extracted document data to markdown."""
    output = []

    # Add metadata if present
    if data['metadata'].get('title'):
        output.append(f"# {data['metadata']['title']}\n")

    # Convert paragraphs
    for para in data['paragraphs']:
        style = para['style']
        text = para['text']

        if style.startswith('Heading'):
            # Extract heading level
            try:
                level = int(style.replace('Heading ', '').replace(' ', ''))
                output.append('#' * level + ' ' + text)
            except ValueError:
                output.append('## ' + text)
        elif style == 'Title':
            output.append('# ' + text)
        elif style == 'Subtitle':
            output.append('## ' + text)
        elif style.startswith('List'):
            output.append('- ' + text)
        else:
            output.append(text)

        output.append('')  # Empty line after paragraph

    # Convert tables
    for table in data['tables']:
        if not table['data']:
            continue

        output.append('')

        # Header row
        header = table['data'][0] if table['data'] else []
        output.append('| ' + ' | '.join(header) + ' |')
        output.append('|' + '|'.join(['---'] * len(header)) + '|')

        # Data rows
        for row in table['data'][1:]:
            # Ensure row has same number of columns as header
            while len(row) < len(header):
                row.append('')
            output.append('| ' + ' | '.join(row[:len(header)]) + ' |')

        output.append('')

    return '\n'.join(output)

def main():
    parser = argparse.ArgumentParser(
        description='Read Word documents and convert to markdown or JSON'
    )
    parser.add_argument('file_path', help='Path to the Word document')
    parser.add_argument('--json', action='store_true', help='Output as JSON')

    args = parser.parse_args()

    try:
        data = read_docx_file(args.file_path)

        if args.json:
            print(json.dumps(data, indent=2, default=str))
        else:
            print(format_as_markdown(data))

    except FileNotFoundError:
        print(f"Error: File not found: {args.file_path}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"Error reading Word document: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
