#!/usr/bin/env python3
"""
Read Word documents (.docx and legacy .doc) and convert to markdown format.

Usage:
    python read_word.py <path_to_word_file> [--json] [--extract-media DIR]

Arguments:
    path_to_word_file    Path to the Word document to read
    --json               Output as JSON instead of markdown (optional)
    --extract-media DIR  Save embedded images (word/media/) to DIR (optional)

Output:
    Markdown text with headings preserved (default) or JSON structure.
    Embedded images are listed (and, with --extract-media, saved); several
    ClinGen supplementary tables ship only as PNGs inside the .docx.

Legacy .doc files (OLE2/CFB format) cannot be read by python-docx. ClinGen
ships some supplementary files in this format, so they are converted first,
using whichever of these is available: LibreOffice (preserves tables),
textutil (macOS), antiword.

Dependencies:
    python-docx (auto-installed if missing)
    Optional for legacy .doc: soffice/libreoffice, textutil, or antiword
"""

import os
import shutil
import subprocess
import sys
import json
import argparse
import tempfile
import zipfile

# OLE2/Compound File Binary signature - marks a legacy .doc (not a zip-based .docx)
OLE2_MAGIC = b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'

def install_dependencies():
    """Install required packages if not present."""
    import subprocess
    try:
        from docx import Document
    except ImportError:
        print("Installing python-docx...", file=sys.stderr)
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "python-docx", "-q"],
            stderr=subprocess.DEVNULL
        )

def is_legacy_doc(file_path):
    """Detect a legacy OLE2 .doc by magic bytes, not by file extension."""
    with open(file_path, 'rb') as f:
        return f.read(8) == OLE2_MAGIC


def read_legacy_doc(file_path, media_out_dir=None):
    """
    Read a legacy .doc by converting it first.

    LibreOffice is tried first because converting to .docx keeps tables intact;
    the text-only converters are fallbacks that lose table structure.
    """
    soffice = shutil.which('soffice') or shutil.which('libreoffice')
    if soffice:
        with tempfile.TemporaryDirectory() as tmpdir:
            try:
                subprocess.run(
                    [soffice, '--headless', '--convert-to', 'docx',
                     '--outdir', tmpdir, file_path],
                    check=True, capture_output=True, timeout=120
                )
                base = os.path.splitext(os.path.basename(file_path))[0] + '.docx'
                converted = os.path.join(tmpdir, base)
                if os.path.exists(converted):
                    print(f"Converted legacy .doc via {os.path.basename(soffice)}",
                          file=sys.stderr)
                    return read_docx_file(converted, media_out_dir)
            except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as e:
                print(f"LibreOffice conversion failed ({e}), trying text extraction",
                      file=sys.stderr)

    # Text-only fallbacks: tables collapse to plain lines
    for cmd, name in ((['textutil', '-convert', 'txt', '-stdout', file_path], 'textutil'),
                      (['antiword', file_path], 'antiword')):
        if not shutil.which(cmd[0]):
            continue
        try:
            out = subprocess.run(cmd, check=True, capture_output=True, timeout=120)
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired):
            continue
        text = out.stdout.decode('utf-8', errors='replace')
        print(f"Converted legacy .doc via {name} (table structure not preserved)",
              file=sys.stderr)
        return {
            'paragraphs': [{'text': line.strip(), 'style': 'Normal'}
                           for line in text.splitlines() if line.strip()],
            'tables': [],
            'images': [],
            'metadata': {}
        }

    raise RuntimeError(
        f"{file_path} is a legacy .doc and no converter is available. "
        "Install LibreOffice (best - keeps tables), or antiword."
    )


def extract_media(docx_path, out_dir=None):
    """
    List (and optionally save) images embedded in a .docx archive.

    python-docx surfaces paragraphs and tables but never the images stored
    under ``word/media/``. Several ClinGen supplementary tables ship only as
    embedded PNGs, so without this step they are invisible to the reader.

    Returns a list of image records ``{name, archive_path, bytes}``; when
    ``out_dir`` is given each image is also written there and gets a
    ``saved_to`` key.
    """
    images = []
    try:
        zf = zipfile.ZipFile(docx_path)
    except (zipfile.BadZipFile, IsADirectoryError, FileNotFoundError):
        return images  # legacy .doc or non-zip: nothing to extract
    with zf:
        names = [n for n in zf.namelist()
                 if n.startswith('word/media/') and not n.endswith('/')]
        if out_dir and names:
            os.makedirs(out_dir, exist_ok=True)
        for n in sorted(names):
            data = zf.read(n)
            rec = {
                'name': os.path.basename(n),
                'archive_path': n,
                'bytes': len(data),
            }
            if out_dir:
                dest = os.path.join(out_dir, os.path.basename(n))
                with open(dest, 'wb') as fh:
                    fh.write(data)
                rec['saved_to'] = dest
            images.append(rec)
    return images


def read_docx_file(file_path, media_out_dir=None):
    """
    Read a .docx file and extract its contents.

    Args:
        file_path: Path to the Word document
        media_out_dir: If set, embedded images are written to this directory

    Returns:
        Dictionary with document structure and content
    """
    try:
        from docx import Document
    except ImportError:
        install_dependencies()
        from docx import Document

    doc = Document(file_path)
    result = {
        'paragraphs': [],
        'tables': [],
        'images': extract_media(file_path, media_out_dir),
        'metadata': {}
    }

    # Extract core properties if available
    try:
        core_props = doc.core_properties
        result['metadata'] = {
            'title': core_props.title or '',
            'author': core_props.author or '',
            'created': str(core_props.created) if core_props.created else '',
            'modified': str(core_props.modified) if core_props.modified else ''
        }
    except Exception:
        pass

    # Extract paragraphs with style information
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style_name = para.style.name if para.style else "Normal"
            result['paragraphs'].append({
                'text': text,
                'style': style_name
            })

    # Extract tables
    for table_idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        result['tables'].append({
            'index': table_idx,
            'rows': len(table.rows),
            'cols': len(table.rows[0].cells) if table.rows else 0,
            'data': table_data
        })

    return result

def format_as_markdown(data):
    """Convert extracted document data to markdown."""
    output = []

    # Add metadata if present
    if data['metadata'].get('title'):
        output.append(f"# {data['metadata']['title']}\n")

    # Convert paragraphs
    for para in data['paragraphs']:
        style = para['style']
        text = para['text']

        if style.startswith('Heading'):
            # Extract heading level
            try:
                level = int(style.replace('Heading ', '').replace(' ', ''))
                output.append('#' * level + ' ' + text)
            except ValueError:
                output.append('## ' + text)
        elif style == 'Title':
            output.append('# ' + text)
        elif style == 'Subtitle':
            output.append('## ' + text)
        elif style.startswith('List'):
            output.append('- ' + text)
        else:
            output.append(text)

        output.append('')  # Empty line after paragraph

    # Convert tables
    for table in data['tables']:
        if not table['data']:
            continue

        output.append('')

        # Header row
        header = table['data'][0] if table['data'] else []
        output.append('| ' + ' | '.join(header) + ' |')
        output.append('|' + '|'.join(['---'] * len(header)) + '|')

        # Data rows
        for row in table['data'][1:]:
            # Ensure row has same number of columns as header
            while len(row) < len(header):
                row.append('')
            output.append('| ' + ' | '.join(row[:len(header)]) + ' |')

        output.append('')

    # Flag embedded images so the reader knows tables-as-PNGs exist
    images = data.get('images') or []
    if images:
        output.append('')
        output.append(f'## Embedded images ({len(images)})')
        output.append('')
        output.append('> These are stored inside the .docx and are NOT rendered '
                      'above. Several ClinGen tables ship only as embedded PNGs; '
                      'extract them with `--extract-media <dir>` to inspect.')
        output.append('')
        for img in images:
            line = f"- `{img['name']}` ({img['bytes']} bytes)"
            if img.get('saved_to'):
                line += f" → saved to `{img['saved_to']}`"
            output.append(line)
        output.append('')

    return '\n'.join(output)

def main():
    parser = argparse.ArgumentParser(
        description='Read Word documents and convert to markdown or JSON'
    )
    parser.add_argument('file_path', help='Path to the Word document')
    parser.add_argument('--json', action='store_true', help='Output as JSON')
    parser.add_argument('--extract-media', metavar='DIR',
                        help='Save embedded images (word/media/) to this directory')

    args = parser.parse_args()

    try:
        if is_legacy_doc(args.file_path):
            data = read_legacy_doc(args.file_path, args.extract_media)
        else:
            data = read_docx_file(args.file_path, args.extract_media)

        if args.json:
            print(json.dumps(data, indent=2, default=str))
        else:
            print(format_as_markdown(data))

    except FileNotFoundError:
        print(f"Error: File not found: {args.file_path}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"Error reading Word document: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
